"""
parsers/resume_text_extractor.py

Day 5 Deliverable: Resume Text Extraction Engine

Converts raw resume files (PDF, DOCX) into clean, normalized, structured text
that downstream AI modules (section segmentation, skill extraction, ATS
scoring, etc.) can rely on.

Public API:
    extract_and_process(file_path, output_dir=None) -> dict
    extract_raw_text(file_path) -> str

Supported inputs: .pdf, .docx
"""

import json
import os
import sys
import time
from pathlib import Path

import docx  # python-docx
import pdfplumber

sys.path.append(str(Path(__file__).resolve().parent.parent))
from utils.text_cleaner import clean_text, normalize_text  # noqa: E402


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class UnsupportedFileTypeError(Exception):
    pass


class ResumeTextExtractor:
    """Extracts and normalizes text from PDF or DOCX resumes."""

    # ---- PDF ----------------------------------------------------------

    def _extract_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF, handling multi-column layouts and tables.

        pdfplumber's default word-order can scramble multi-column resumes,
        so we sort words by (top, x0) per page as a lightweight column-aware
        reading order, and separately pull out any detected tables (common
        for 'Skills' grids or two-column layouts) as their own text block.
        """
        page_texts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                words = page.extract_words(use_text_flow=False, keep_blank_chars=False)
                if words:
                    # Group words into lines by rounded 'top' position, then
                    # sort each line left-to-right. This keeps two-column
                    # resumes from interleaving mid-sentence.
                    words_sorted = sorted(words, key=lambda w: (round(w["top"] / 3), w["x0"]))
                    lines = []
                    current_line, current_top = [], None
                    for w in words_sorted:
                        top_bucket = round(w["top"] / 3)
                        if current_top is None or top_bucket == current_top:
                            current_line.append(w)
                            current_top = top_bucket
                        else:
                            lines.append(" ".join(x["text"] for x in sorted(current_line, key=lambda x: x["x0"])))
                            current_line = [w]
                            current_top = top_bucket
                    if current_line:
                        lines.append(" ".join(x["text"] for x in sorted(current_line, key=lambda x: x["x0"])))
                    page_texts.append("\n".join(lines))
                else:
                    # Fallback for pages where word extraction fails
                    fallback = page.extract_text() or ""
                    page_texts.append(fallback)

                # Pull tables separately and append as labeled blocks so
                # tabular skill/qualification grids aren't lost or garbled.
                tables = page.extract_tables()
                for t_idx, table in enumerate(tables):
                    rows_text = []
                    for row in table:
                        cells = [c.strip() for c in row if c and c.strip()]
                        if cells:
                            rows_text.append(" | ".join(cells))
                    if rows_text:
                        page_texts.append("[TABLE]\n" + "\n".join(rows_text))

        return "\n\n".join(t for t in page_texts if t and t.strip())

    # ---- DOCX -----------------------------------------------------------

    def _extract_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX, preserving paragraph order and pulling
        tables (frequently used for skills matrices / two-column layouts)
        as labeled blocks, in their position within the document body.
        """
        document = docx.Document(file_path)
        blocks = []

        # python-docx doesn't expose a single ordered stream of paragraphs +
        # tables out of the box, so we walk the underlying XML body in order.
        body = document.element.body
        for child in body.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "p":
                para = next((p for p in document.paragraphs if p._p is child), None)
                if para is not None and para.text.strip():
                    blocks.append(para.text)
            elif tag == "tbl":
                table = next((t for t in document.tables if t._tbl is child), None)
                if table is not None:
                    rows_text = []
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if cells:
                            rows_text.append(" | ".join(cells))
                    if rows_text:
                        blocks.append("[TABLE]\n" + "\n".join(rows_text))

        return "\n".join(blocks)

    # ---- Dispatcher -------------------------------------------------------

    def extract_raw_text(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self._extract_pdf(file_path)
        elif ext == ".docx":
            return self._extract_docx(file_path)
        else:
            raise UnsupportedFileTypeError(
                f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
            )

    def extract_and_process(self, file_path: str, output_dir: str = None) -> dict:
        """
        Full pipeline: extract -> clean -> normalize -> (optionally) persist.

        Returns a structured dict:
            {
              "source_file": str,
              "file_type": str,
              "char_count": int,
              "extracted_at": float (epoch),
              "raw_text": str,
              "cleaned_text": str,
              "output_path": str | None
            }
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(file_path)

        raw = self.extract_raw_text(file_path)
        cleaned = clean_text(raw)
        normalized = normalize_text(cleaned)

        result = {
            "source_file": os.path.basename(file_path),
            "file_type": Path(file_path).suffix.lower().lstrip("."),
            "char_count": len(normalized),
            "extracted_at": time.time(),
            "raw_text": raw,
            "cleaned_text": normalized,
            "output_path": None,
        }

        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            stem = Path(file_path).stem
            out_path = os.path.join(output_dir, f"{stem}.json")
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(result, f, indent=2, ensure_ascii=False)
            # Also drop a plain .txt for quick human inspection
            txt_path = os.path.join(output_dir, f"{stem}.txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(normalized)
            result["output_path"] = out_path

        return result


def extract_and_process(file_path: str, output_dir: str = None) -> dict:
    """Module-level convenience wrapper (functional API)."""
    return ResumeTextExtractor().extract_and_process(file_path, output_dir)


if __name__ == "__main__":
    # Simple CLI usage: python resume_text_extractor.py <file> [output_dir]
    if len(sys.argv) < 2:
        print("Usage: python resume_text_extractor.py <resume_file> [output_dir]")
        sys.exit(1)
    target = sys.argv[1]
    out_dir = sys.argv[2] if len(sys.argv) > 2 else "outputs/extracted"
    res = extract_and_process(target, out_dir)
    print(f"Extracted {res['char_count']} chars -> {res['output_path']}")

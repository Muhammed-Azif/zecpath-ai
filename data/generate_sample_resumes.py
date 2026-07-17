"""
data/generate_sample_resumes.py

Generates sample resumes (PDF + DOCX) with realistic noisy formatting -
bullet glyph variety, a skills table, mixed capitalization headings - so the
extraction engine has real fixtures to be tested against.

Run: python generate_sample_resumes.py
Outputs into: data/sample_resumes/
"""

import os

import docx
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import Table, TableStyle

OUT_DIR = os.path.join(os.path.dirname(__file__), "sample_resumes")
os.makedirs(OUT_DIR, exist_ok=True)


def make_pdf_resume(path):
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    y = height - 60

    def line(text, size=11, dy=16, bold=False):
        nonlocal y
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        c.drawString(60, y, text)
        y -= dy

    line("PRIYA SHARMA", 16, bold=True)
    line("priya.sharma@example.com | +91-9876543210 | Bengaluru, India", 10)
    y -= 10

    line("SUMMARY", 12, bold=True)
    line("Backend engineer with 4 years of experience building scalable APIs", 10)
    line("and event-driven systems using Python and Node.js.", 10)
    y -= 10

    line("SKILLS", 12, bold=True)
    line("- Python  - Django  - FastAPI  - PostgreSQL", 10)
    line("- Docker  - AWS (EC2, S3, Lambda)  - Kafka", 10)
    y -= 10

    line("WORK EXPERIENCE", 12, bold=True)
    line("Backend Engineer, Nimbus Systems (2021 - Present)", 11, bold=True)
    line("- Designed a microservices architecture serving 2M+ daily requests", 10)
    line("- Reduced API latency by 38% through query optimization", 10)
    y -= 6
    line("Software Engineer, CodeForge Labs (2019 - 2021)", 11, bold=True)
    line("* Built internal tooling for automated deployment pipelines", 10)
    y -= 10

    line("EDUCATION", 12, bold=True)
    line("B.Tech in Computer Science, RV College of Engineering, 2019", 10)

    # A genuine bordered table (not just column-positioned text) so
    # pdfplumber's grid-based table detection has real lines to find.
    y -= 20
    c.setFont("Helvetica-Bold", 12)
    c.drawString(60, y, "CERTIFICATIONS")
    y -= 10

    table_data = [
        ["Certification", "Year"],
        ["AWS Certified Developer", "2022"],
        ["Certified Kubernetes Administrator", "2023"],
    ]
    table = Table(table_data, colWidths=[280, 80])
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.75, colors.black),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
    ]))
    table_width, table_height = table.wrapOn(c, 400, 200)
    table.drawOn(c, 60, y - table_height)

    c.save()


def make_docx_resume(path):
    d = docx.Document()

    h = d.add_paragraph()
    run = h.add_run("RAHUL VERMA")
    run.bold = True
    run.font.size = Pt(18)

    d.add_paragraph("rahul.verma@example.com | +91-9123456780 | Pune, India")
    d.add_paragraph("")

    d.add_paragraph("OBJECTIVE").runs[0].bold = True
    d.add_paragraph(
        "Frontend developer with 3 years of experience in React and TypeScript, "
        "seeking a role building accessible, high-performance web applications."
    )

    d.add_paragraph("")
    p = d.add_paragraph("skills")  # deliberately lowercase to test heading normalization
    p.runs[0].bold = True
    d.add_paragraph("\u2022 React  \u2022 TypeScript  \u2022 Redux  \u2022 Tailwind CSS  \u2022 Jest")

    d.add_paragraph("")
    d.add_paragraph("Work Experience").runs[0].bold = True
    d.add_paragraph("Frontend Developer, Bluepeak Digital (2022 - Present)").runs[0].bold = True
    d.add_paragraph("\u2022 Rebuilt the design system used across 6 internal products")
    d.add_paragraph("\u2022 Improved Lighthouse performance score from 61 to 94")

    d.add_paragraph("")
    d.add_paragraph("EDUCATION").runs[0].bold = True
    d.add_paragraph("B.E. in Information Technology, Pune Institute of Technology, 2021")

    # Skills matrix as an actual Word table, to test table extraction
    d.add_paragraph("")
    d.add_paragraph("CERTIFICATIONS").runs[0].bold = True
    table = d.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Certification"
    hdr[1].text = "Year"
    for cert, year in [("Meta Front-End Developer Certificate", "2023"),
                        ("Google UX Design Certificate", "2022")]:
        row = table.add_row().cells
        row[0].text = cert
        row[1].text = year

    d.save(path)


if __name__ == "__main__":
    pdf_path = os.path.join(OUT_DIR, "priya_sharma_resume.pdf")
    docx_path = os.path.join(OUT_DIR, "rahul_verma_resume.docx")
    make_pdf_resume(pdf_path)
    make_docx_resume(docx_path)
    print(f"Created:\n  {pdf_path}\n  {docx_path}")
